# pyright: reportMissingImports=false

import json
import random
import re
import sqlite3
import requests
import wikipedia
from datetime import datetime
from docx import Document
from urllib.request import urlopen
from revChatGPT.V1 import Chatbot


#web

import pyrebase
import firebase_admin
from firebase_admin import auth as auth1
from firebase_admin import credentials, firestore

#kivy

from kivy.uix.screenmanager import SlideTransition
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.image import Image
from kivy.uix.button import Button

from kivy.core.text import LabelBase
from kivy.core.clipboard import Clipboard
from kivy.clock import Clock
from kivy.properties import  ListProperty, ObjectProperty,StringProperty, NumericProperty
from kivy.core.window import Window
from kivy.core.text import LabelBase
from kivy.metrics import dp
from kivy.uix.image import Image

#kivymd

from kivymd.uix.dialog import MDDialog
from kivymd.uix.floatlayout import MDFloatLayout
from kivymd.uix.list import ILeftBodyTouch, TwoLineAvatarIconListItem
from kivymd.uix.pickers import MDDatePicker
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.selectioncontrol import MDCheckbox
from kivymd.uix.card import MDCard
from kivymd.uix.dialog import MDDialog
from kivymd.uix.button import MDFlatButton
from kivymd.uix.label import MDLabel

from kivymd.app import MDApp
from kivymd.toast import toast

#Permissions

global internet_connection
internet_connection = False
try:
    response = requests.get("http://www.google.com")
    if response.status_code == 200:
        internet_connection = True
except requests.ConnectionError:
    pass  

#Internet Libraries


cred = credentials.Certificate({
    "type": "service_account",
    "project_id": "simplifique-c881d",
    "private_key_id": "83f7f7868abc181d033a4d5e2c8ffe1c8af26688",
    "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvQIBADANBgkqhkiG9w0BAQEFAASCBKcwggSjAgEAAoIBAQDTi/op1fBF7Sx1\nv4+XkScBR42e31isc0olzMB/+ypswa4W7DmDywAxn30TASjWtCQeTdIBeiUeJstc\n7QGXku165kFnrUq6QdyOf6O/aGJibKHpuwY91NiFEO5/XkxwgmKbhUnjlWuNe8f1\nRB+o0Jiv833T8nfyiygtSaaDLgzlw0/lZ47NNfC9/kXprZrRFMDUhxlSB8wpHNQp\nbzQMLryYQx8/NfC4xr6NWI+v+aO32v65wGaXqpRpW7Tq81YBnXccND68EqKBJNTi\nfNIjmhbXj+DzFtgavK4oXPtHm4SSNq0q85jzc1e9oW6yS7gVpscrsXstOOdeunhm\nhTJ00qJnAgMBAAECggEAWTMirtVI1RNmtdeqQmywF7gGHUFr8HtEfp/RY6WSg/0+\n3OeKcOn+EO6BHKxWfgHDYQvLS1gnookVIv/Ethb8D+BbH85QPi1bRLLJZwIqyfmo\nZBe6UAQZsDItfoNSk/ZGgfE38MCmcygIboDlIJekajyvh9krfpfyvvXZQmL+iTqH\nZuqkqHLubfFaKEvL9jcrcYm1Z0ZLoGx4h9rLjc/0X8MMndkhuZYYn0CCn4i5fop8\nwxxsuACcIk1ug7ZiVJb6qMN2FSMEO7csghoCPHl0dP2RaaGpg2ILsDQhgGIWjRw3\nhdBj5sLXDdD/ZHlSburs5PjzxYPls9GRUfGQ8i5GMQKBgQDxtPwBob0uqM338LZN\niEcAMP5V70hGU8ryLXRRz3SCKmMxKJPeFsajVchfrNxexIQ/5MuWBSm9BO/g/qTN\nJDeFaQVPRIgTtT0xITNHcMUdT8rBwB32I7SwSkryIxsLUnaUAqA11WurUQk0cURb\nuFBZ3f6fDbQonl4oWwK2wR3+4wKBgQDgDmvHGLAve19A26fCXLE0Q87065RYLILh\nVTK2umNi4NPA8jWmxncI9F7eEBcEEph4aTlemmvBX2d3WNp5t/DyiTwszqrzTbeA\nikqrOqTxGJV24WvDehVI5zTVdqiKc0my+dJOUX5dbwV09I63nL77TWKXELfCemfw\nCbRL5v+BrQKBgGbLceG/x5VwdShdVyriKlAKhiBGA5blTApzCmVAtWwmWsktWLW7\nOf99HBqUiaREL3p885h52aZp0xr9MVmNbY9verKbksPO8JdUZ1qauzocFT8RVay4\nwr+22OjhxT6rc4K/GyPKAGB7tk53XXskiAewQfmi+lvL/n9rNVxEBV3BAoGAeEXF\nfT63dQWZAEvpJeB0D0ZHFhpPq3VZXHRLoOM07qMZiH18Z2YqB9iGBFZGxJzm09xI\nO4xRQ6Be/iXoQWaIJOmeL79Q7QJO+uVBZ+E3IWS89u/S1T/3pQbXya7Ekm2IplaM\nmhYM60Lpfvq4kb/GlUfZIJaMzgy/No8/BW+ewJECgYEArd8O6vFKY5y26zaMy6zN\nVi/3SdN2mXs96a0N6h+qhTcly53+P9Xpa/8xOZxypc26gf38JB/7VUaWX7FYDz20\n5GqleRXVSxmosaL9Hib9zGyaVlMlV/ytgIc150+C6SXJg5onwTZ+DKCuY7kBwIk6\nrx38QKyJd/+p5XbmTJcFPUo=\n-----END PRIVATE KEY-----\n",
    "client_email": "firebase-adminsdk-arl0m@simplifique-c881d.iam.gserviceaccount.com",
    "client_id": "110928610409474987418",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/firebase-adminsdk-arl0m%40simplifique-c881d.iam.gserviceaccount.com"
    })
firebase_admin.initialize_app(cred)


class Accont:
    def __init__(self):
        firebaseConfig = {
            "apiKey": "AIzaSyCC_WAxsEv_GA99LBhucJ33_cWWy2aENCo",
            "authDomain": "simplifique-c881d.firebaseapp.com",
            "projectId": "simplifique-c881d",
            "databaseURL": "https://simplifique-c881d-default-rtdb.firebaseio.com",
            "storageBucket": "simplifique-c881d.appspot.com",
            "messagingSenderId": "413832219747",
            "appId": "1:413832219747:web:8108fd5f8557737180bc56",
            "measurementId": "G-2M4RF2PBPP"
        }
        firebase = pyrebase.initialize_app(firebaseConfig)
        self.auth = firebase.auth()

    def sign_in(self, email, password):
        if internet_connection:
            try:
                self.auth.sign_in_with_email_and_password(email, password)
                return True

            except:
                return False
        else:
            toast('No Internet!')
    def sing_up(self, email, password, name):
        if internet_connection:
            try:
                auth1.create_user(email = email, email_verified = False, password = password, display_name = name, disabled = False)
                email = email
                user = auth1.get_user_by_email(email)
                user_id = user.uid
                display_name = user.display_name
                data = {'name': str(display_name), 'email': str(email)}
                db1 = firestore.client()
                db1.collection('users').document(user_id).set(data)
                return True
            except:
                return False
        else:
            toast("No Internet")
    def reset_password(self, email):
        if internet_connection:
            try:
                self.auth.send_password_reset_email(email)
                return True

            except:
                return False
        else:
            toast("No Internet")
class Database:
    def __init__(self):
        self.con = sqlite3.connect('data.db')
        self.cursor = self.con.cursor()
        self.create_task_table()

    def create_task_table(self):
        """Create tasks table"""
        self.cursor.execute("CREATE TABLE IF NOT EXISTS tasks(id integer PRIMARY KEY AUTOINCREMENT, task varchar(50) NOT NULL, due_date varchar(50), completed BOOLEAN NOT NULL CHECK (completed IN (0, 1)))")
        self.con.commit()
        

    def create_task(self, task, due_date=None):
        """Create a task"""
        self.cursor.execute("INSERT INTO tasks(task, due_date, completed) VALUES(?, ?, ?)", (task, due_date, 0))
        self.con.commit()

        # GETTING THE LAST ENTERED ITEM SO WE CAN ADD IT TO THE TASK LIST
        created_task = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE task = ? and completed = 0", (task,)).fetchall()
        return created_task[-1]

    def get_tasks(self):
        """Get tasks"""
        uncomplete_tasks = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE completed = 0").fetchall()
        completed_tasks = self.cursor.execute("SELECT id, task, due_date FROM tasks WHERE completed = 1").fetchall()

        return completed_tasks, uncomplete_tasks

    

    def mark_task_as_complete(self, taskid):
        """Marking tasks as complete"""
        self.cursor.execute("UPDATE tasks SET completed=1 WHERE id=?", (taskid,))
        self.con.commit()

    def mark_task_as_incomplete(self, taskid):
        """Mark task as uncomplete"""
        self.cursor.execute("UPDATE tasks SET completed=0 WHERE id=?", (taskid,))
        self.con.commit()

        # return the text of the task
        task_text = self.cursor.execute("SELECT task FROM tasks WHERE id=?", (taskid,)).fetchall()
        return task_text[0][0]

    def delete_task(self, taskid):
        """Delete a task"""
        self.cursor.execute("DELETE FROM tasks WHERE id=?", (taskid,))
        self.con.commit()

    def close_db_connection(self):
        self.con.close()

#Screens

class Screen1(Screen):
    def on_pre_enter(self, *args):
        try:
            conn = sqlite3.connect('data.db')
            cursor = conn.cursor()
            cursor.execute('SELECT email, password FROM users')
            result = cursor.fetchone()
            conn.close()
            if result:
                email, password = result
                self.ids.l_email.text = email
                self.ids.l_password.text = password
            else:
                self.ids.l_email.text = ""
                self.ids.l_password.text = ""
        except sqlite3.Error:
            self.ids.l_email.text = ""
            self.ids.l_password.text = ""
        except Exception:
            self.ids.l_email.text = ""
            self.ids.l_password.text = ""
        else:
            pass

    def try_sign_in(self):
        Clock.schedule_once(self.start_sign_in, .5)

    def start_sign_in(self, *args):
        email = self.ids.l_email.text
        password = self.ids.l_password.text

        if accont.sign_in(email, password):
            thing = MDApp.get_running_app()
            thing.root.current = "screen4"
            user = auth1.get_user_by_email(email)
            user_id = user.uid
            name = user.display_name
            conn = sqlite3.connect('data.db')
            cursor = conn.cursor()
            cursor.execute('DROP TABLE IF EXISTS users')
            cursor.execute('''CREATE TABLE IF NOT EXISTS users (email TEXT, password TEXT, name TEXT, uid TEXT)''')
            cursor.execute('INSERT OR REPLACE INTO users VALUES (?, ?, ?, ?)', (email, password, name, user_id))
            conn.commit()
            conn.close()
            thing.press()
            toast("Logged In Successfully!")
        else:
            if email and password:
                self.ids.l_password.text = ""
                self.ids.l_email.focus = True

            elif email and not password:
                self.ids.l_password.focus = True
                toast("This password isn't correct.")

            elif password and not email:
                self.ids.l_password.text = ""
                self.ids.l_email.focus = True
                toast("This email isn't in our database.")

            else:
                self.ids.l_email.focus = True
class Screen2(Screen):
    def on_pre_enter(self, *args):
        self.ids.s_email.text = ""
        self.ids.s_password.text = ""
        self.ids.s_name.text = ""

    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"

    def validate_info(self):
        email = self.ids.s_email.text
        password = self.ids.s_password.text
        name = self.ids.s_name.text

        if password == password and "@" in email and ".com" in email and len(password) >= 6:
            if accont.sing_up(email, password, name):
                toast("Registration Created Successfully!")
                Clock.schedule_once(self.callback, 3)
            else:
                toast("Failed to create record.")

        else:
            if not email or "@" not in email or ".com" not in email:
                self.ids.s_email.focus = True

            elif len(password) < 6:
                self.ids.s_name.text = ""
                self.ids.s_password.text = ""
                self.ids.s_password.focus = True
class Screen3(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"
    
    def on_pre_enter(self, *args):
        self.ids.r_email.text = ""

    def send_email_confirm(self):
        email = self.ids.r_email.text

        if accont.reset_password(email):
            toast("Email Succesfully Sended!")
        else:
            toast("Failed to send email.")

        Clock.schedule_once(self.callback, 3)
class Screen4(Screen):
    def callback(self, *args):
        MDApp.get_running_app().root.current = "screen1"
class Screen5(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen6(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen7(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
    def essay_helper(self, *args):
        w_name = self.ids.wiki_name.text
        language = self.ids.wiki_language.text
        title = self.ids.wiki_title.text
        wikipedia.set_lang(language)
        try:
            if title:
                wiki = wikipedia.page(title)
            else:
                self.ids.wiki_title.text = ""
                self.ids.wiki_title.focus = True
                toast("Project name invalid")
                return
        except wikipedia.exceptions.PageError:
            self.ids.wiki_title.text = ""
            self.ids.wiki_title.focus = True
            toast(" Project name not found \n Please enter another title ")
            return
        text = wiki.content
        text = re.sub(r'==', '', text)
        text = re.sub(r'=', '', text)
        text = re.sub(r'\n', '\n    ', text)
        split = text.split('See also', 1)
        text = split[0] 
        document = Document()
        paragraph = document.add_heading(title, 0)
        paragraph.alignment = 1
        paragraph = document.add_paragraph('    ' + text)
        paragraph = document.add_paragraph(w_name)
        paragraph.alignment = 2
        document.save(title + ".docx")
        toast(" Project saved in local directory ")
class Screen8(Screen):
    global internet_connection
    image_source = StringProperty()
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
    def  meal_text(self, *args):
        if internet_connection:
            try:
                url = 'https://www.themealdb.com/api/json/v1/1/random.php'
                meal = requests.get(url).json()
                meals = meal["meals"]
                ingredients = []
                measures = []
                links = []
                meal_ingredients = ""
                global image_link
                for i in range(20):
                    for ingredient in meals:
                        if 1==1:
                            ingredients.append(ingredient[f"strIngredient{i+1}"])
                        else:
                            break
                    for measure in meals:
                        if 1==1:
                            measures.append(measure[f"strMeasure{i+1}"])
                        else:
                            break
                for j in range(20):
                    try: 
                        meal_ingredients = meal_ingredients + "• " + ingredients[j] + " - " + measures[j] + "\n"
                    except:
                        toast('Try again')
                        pass
                for link in meals:
                    links.append(link)
                    image_link = link["strMealThumb"]
                    yt_link = link["strYoutube"]
                    meal_name = link['strMeal']
                    meal_type = link['strCategory']
                    meal_area = link['strArea']
                    meal_instructions = link['strInstructions']
                try:
                    meal_description = meal_name +'\n\n' + 'Category: ' + meal_type +'\n' + 'Area: ' + meal_area +'\n\n' + 'Instructions: \n\n' + meal_instructions +'\n'
                    meal_ingredients_text = "Ingredients: \n\n" + meal_ingredients +  '\n' + "•  -" 
                    meal_text = meal_description +'\n' + meal_ingredients_text.replace("•  -", "")
                    self.ids.meal_text.text = meal_text
                    self.image_source = image_link
                except:
                    toast('Try Again')
                    pass
                try:
                    Clipboard.copy(yt_link)
                    MDApp.get_running_app().open_url(yt_link)
                except:
                    toast('Try Again')
                    pass
                toast("Meal Found")
            except :
                toast("An error occurred")
        else:
            toast("No internet connection - R")
class Screen9(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen10(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen11(Screen):
    def __init__(self, **kwargs):
        super(Screen11, self).__init__(**kwargs)
        self.description = kwargs.get('description', '')
        self.url_link = kwargs.get('url_link', '')
        global internet_connection
        if internet_connection:
            try:
                art = "school"
                api_key = '283533136981441da324ba7c1b5d0cc5'
                url = f'https://newsapi.org/v2/everything?q={art}&apikey='+api_key
                news = requests.get(url).json()
                articles = news["articles"]
                for i in range(10): # Print the first five articles
                    url_link = articles[i]["url"]
                    description = articles[i]["description"]
                    words = description.split()[:30]
                    description = ' '.join(words) + '...'
                    img_link = articles[i]["urlToImage"]
                    card = NewsCard(source=img_link, description=description, url_link=url_link,padding=15, radius=15,elevation=3)
                    self.ids.news_list.add_widget(card)
                toast("Reload Succesfull")
            except Exception as e:
                toast("An error occurred")
                print(e)
        else: 
            toast("No internet connection - N")
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen12(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen10"
class Screen13(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen14(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen15(Screen):
    def callback(self, *args):
        app = MDApp.get_running_app()
        app.root.transition = SlideTransition(direction="right")
        app.root.current = "screen4"
class Screen16(Screen):  
    def __init__(self, **kwargs):
        super(Screen16, self).__init__(**kwargs)
        note =  """I recently came across an incredibly beautiful UI app that caught my attention. The app features a sleek and modern design with a seamless user experience.

The color scheme of the app is stunning, using a combination of vibrant and complementary colors that create a visually pleasing interface. The use of gradients and subtle shadows adds depth and dimension to the UI elements, making them appear more engaging and interactive.

The typography in the app is carefully chosen, with elegant and legible fonts that enhance the overall aesthetic. The font sizes and styles are consistent throughout the app, creating a sense of harmony and coherence.

The layout and organization of the app are intuitive and user-friendly. The navigation is smooth, and the placement of buttons and interactive elements is thoughtfully designed for easy accessibility. The app's content is well-structured, making it effortless to find information and navigate through different sections.

Furthermore, the app incorporates stunning illustrations, icons, and visual elements that contribute to its overall beauty. The use of animations and transitions adds a touch of dynamism and liveliness, making the app feel dynamic and engaging.

Overall, this beautiful UI app sets a high standard for aesthetic design and user experience. It exemplifies the power of thoughtful design choices and attention to detail, creating an enjoyable and visually captivating experience for its users.
        """
        self.ids.note_text.text = note     
    def callback(self, *args):
            app = MDApp.get_running_app()
            app.root.transition = SlideTransition(direction="right")
            app.root.current = "screen9"

#Functions

class DialogContent(MDBoxLayout):
    """OPENS A DIALOG BOX THAT GETS THE TASK FROM THE USER"""
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        # set the date_text label to today's date when useer first opens dialog box
        self.ids.date_text.text = str(datetime.now().strftime('%A %d %B %Y'))


    def show_date_picker(self):
        """Opens the date picker"""
        date_dialog = MDDatePicker()
        date_dialog.bind(on_save=self.on_save)
        date_dialog.open()

    def on_save(self, instance, value, date_range):
        """This functions gets the date from the date picker and converts its it a
        more friendly form then changes the date label on the dialog to that"""

        date = value.strftime('%A %d %B %Y')
        self.ids.date_text.text = str(date)
class ListItemWithCheckbox(TwoLineAvatarIconListItem):
    '''Custom list item'''

    def __init__(self, pk=None, **kwargs):
        super().__init__(**kwargs)
        # state a pk which we shall use link the list items with the database primary keys
        self.pk = pk


    def mark(self, check, the_list_item):
        '''mark the task as complete or incomplete'''
        if check.active == True:
            the_list_item.text = '[s]'+the_list_item.text+'[/s]'
            db.mark_task_as_complete(the_list_item.pk)# here
        else:
            the_list_item.text = str(db.mark_task_as_incomplete(the_list_item.pk))# Here

    def delete_item(self, the_list_item):
        '''Delete the task'''
        self.parent.remove_widget(the_list_item)
        db.delete_task(the_list_item.pk)# Here
class LeftCheckbox(ILeftBodyTouch, MDCheckbox):
    '''Custom left container'''
class ProfileCard(MDFloatLayout):
    pass
class Command(MDLabel):
    text = StringProperty()
    size_hint_x = NumericProperty()
    halign = StringProperty()
    font_name = "Poppins"
    font_size = 17
class Response(MDLabel):
    text = StringProperty()
    size_hint_x = NumericProperty()
    halign = StringProperty()
    font_name = "Poppins"
    font_size = 17
class ResponseImage(Image):
    source = StringProperty()
class NewsCard(MDCard):
    description = StringProperty()
    source = StringProperty()
    url_link = StringProperty()
class ScreenSwitcher:
    def switch_screen1(self, *args):
        self.root.transition = SlideTransition(direction="right")
        self.root.current = "screen1"

    def switch_screen2(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen2"

    def switch_screen3(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen3"

    def switch_screen4(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen4"
    
    def switch_screen5(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen5"
    
    def switch_screen6(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen6"
    
    def switch_screen7(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen7"
    
    def switch_screen8(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen8"
    
    def switch_screen9(self, *args):
        self.root.transition = SlideTransition(direction="left") #FadeTransition(duration=0.5)
        self.root.current = "screen9"
    
    def switch_screen10(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen10"

    def switch_screen11(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen11"

    def switch_screen12(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen12"

    def switch_screen13(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen13"

    def switch_screen14(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen14"

    def switch_screen15(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen15"

    def switch_screen16(self, *args):
        self.root.transition = SlideTransition(direction="left")
        self.root.current = "screen16"
class RateUs(BoxLayout):
    def select1(self, *args):
        self.ids.star1.opacity = 1
        self.ids.star2.opacity = 0.15
        self.ids.star3.opacity = 0.15
        self.ids.star4.opacity = 0.15
        self.ids.star5.opacity = 0.15
    def select2(self, *args):
        self.ids.star1.opacity = 1
        self.ids.star2.opacity = 1
        self.ids.star3.opacity = 0.15
        self.ids.star4.opacity = 0.15
        self.ids.star5.opacity = 0.15
    def select3(self, *args):
        self.ids.star1.opacity = 1
        self.ids.star2.opacity = 1
        self.ids.star3.opacity = 1
        self.ids.star4.opacity = 0.15
        self.ids.star5.opacity = 0.15
    def select4(self, *args):
        self.ids.star1.opacity = 1
        self.ids.star2.opacity = 1
        self.ids.star3.opacity = 1
        self.ids.star4.opacity = 1
        self.ids.star5.opacity = 0.15
    def select5(self, *args):
        self.ids.star1.opacity = 1
        self.ids.star2.opacity = 1
        self.ids.star3.opacity = 1
        self.ids.star4.opacity = 1
        self.ids.star5.opacity = 1
    def send(self):
        toast("Rated")
class Feedback(BoxLayout):
    def send_email(self):
        name = "Marius"
        toast(f"Hi, {name}, thanks for feedback!")
class Report(BoxLayout):
    def send(self):
        toast("Thanks for your contribution!")
class OptionButton(Button):
    bg_color = ListProperty([1,1,1,1])

#Main 

class MainApp(MDApp, ScreenManager, BoxLayout, Screen10, ScreenSwitcher):
    global sm, chatbot
    api_key = "60aae825eb1705b59a97532605cbae66"
    chatbot = Chatbot(config={"access_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCIsImtpZCI6Ik1UaEVOVUpHTkVNMVFURTRNMEZCTWpkQ05UZzVNRFUxUlRVd1FVSkRNRU13UmtGRVFrRXpSZyJ9.eyJodHRwczovL2FwaS5vcGVuYWkuY29tL3Byb2ZpbGUiOnsiZW1haWwiOiJtYXJpdXMuZ2FicnllbDIwMTdAZ21haWwuY29tIiwiZW1haWxfdmVyaWZpZWQiOnRydWV9LCJodHRwczovL2FwaS5vcGVuYWkuY29tL2F1dGgiOnsidXNlcl9pZCI6InVzZXItYkhhenpoOTlCZHlXQ1RNbFdEdWFpWEE0In0sImlzcyI6Imh0dHBzOi8vYXV0aDAub3BlbmFpLmNvbS8iLCJzdWIiOiJnb29nbGUtb2F1dGgyfDEwNDY2NzIwNTAwMjA2MjA3NjY3MiIsImF1ZCI6WyJodHRwczovL2FwaS5vcGVuYWkuY29tL3YxIiwiaHR0cHM6Ly9vcGVuYWkub3BlbmFpLmF1dGgwYXBwLmNvbS91c2VyaW5mbyJdLCJpYXQiOjE2ODU3MDg0MzksImV4cCI6MTY4NjkxODAzOSwiYXpwIjoiVGRKSWNiZTE2V29USHROOTVueXl3aDVFNHlPbzZJdEciLCJzY29wZSI6Im9wZW5pZCBwcm9maWxlIGVtYWlsIG1vZGVsLnJlYWQgbW9kZWwucmVxdWVzdCBvcmdhbml6YXRpb24ucmVhZCBvcmdhbml6YXRpb24ud3JpdGUifQ.IhSC4jCp70082le4Cw-nqJBwgfgBDpqls6qc7dwj2YzcGUMeLFgCMdGZg1aei5HjRztH9xpoN8kqgVoFVyxmlR6V4uegwt7NBcPeWXDJK8fHUHDImGeQMbVbuJH72MSW6e39L-yk9JujHVRF_B2Gv_Yfa6KQzFK6blk3AKlVcq87Ko9m-OJahC0NzWfbE8gf6l9zNCuGL0hBa9LJqGch4XhZQMpUxAzMNIMJravc_nsSC-5WpTo45gt6dpKQQ-lIznRVxY-5PxIVhKWVpZkkOXVzSnNRnqtakjGXw8e6Ii1YcMtEJEK9lZf2ZelbF6GWxHhV6sEF1BX6nwYw1CV9MA"})
    sm = ScreenManager()
    select_sign = ""
    answer = ""
    correct = 0
    wrong = 0
    task_list_dialog = r_dialog = f_dialog = None
    dropdown = ObjectProperty
    def build(self):
        self.icon = "images/icon.png"
        self.title='Simplifique'
        self.theme_cls.theme_style = "Light"
        self.theme_cls.primary_palette = "DeepPurple"
        self.theme_cls.primary_hue = "500"
        
        #Internet Screen
        if internet_connection:
            sm.add_widget(Screen1(name='screen1')) # Login Page
            sm.add_widget(Screen2(name='screen2')) # SignUp Page
            sm.add_widget(Screen3(name='screen3')) # Reset Password
        
        #Window settings
        width, height = 450, 1000 # 9*20
        Window.size = (width, height)
        Window.left = 1  
        Window.top = 30

        #Graphic Screen
        sm.add_widget(Screen4(name='screen4')) # Home Screen
        sm.add_widget(Screen5(name='screen5')) # ToDo 
        sm.add_widget(Screen6(name='screen6')) # Weather 
        sm.add_widget(Screen7(name='screen7')) # Essay Helper
        sm.add_widget(Screen8(name='screen8')) # Meal 
        sm.add_widget(Screen9(name='screen9')) # Contact 
        sm.add_widget(Screen10(name='screen10')) # Math
        sm.add_widget(Screen11(name='screen11')) # News
        sm.add_widget(Screen12(name='quiz')) # 
        sm.add_widget(Screen13(name='screen13')) # Setings
        sm.add_widget(Screen14(name='screen14')) # Image
        sm.add_widget(Screen15(name='screen15')) # Chat 
        sm.add_widget(Screen16(name='screen16'))
        return sm
    def select_sign(self, sign):
        self.selected_sign = sign
        num1 = random.randint(1,10)
        num2 = random.randint(1,10)
        sm.get_screen("quiz").ids.question.text = f"{num1} {sign} {num2} = ?"
        if sign == "+":
            self.answer = str(num1+num2)
        elif sign == "-":
            self.answer = str(num1-num2)
        elif sign == "÷":
            self.answer = str(round((num1/num2),1))
        elif sign == "×":
            self.answer = str(num1*num2)
        option_list = [self.answer]
        option_len = 1
        while option_len < 4:
            option = 0
            if sign == "+":
                option = str(random.randint(1,20))
            elif sign == "-":
                option = str(random.randint(-10,10))
            elif sign == "÷":
                option = str(round(random.uniform(1,20),1))
            elif sign == "×":
                option = str(random.randint(1,100))
            if option not in option_list:
                option_list.append(option)
            else:
                option_len -= 1
            option_len  += 1
        random.shuffle(option_list)
        for i in range(1, 5):
            sm.get_screen("quiz").ids[f"option{i}"].text = str(option_list[i-1])
        sm.current = "quiz"
    def get_id(self,instance):
        for id, widget in instance.parent.parent.parent.ids.items():
            if widget.__self__ == instance:
                return id
    def quiz(self,option,instance):
        if option == self.answer:
            self.correct += 1
            sm.get_screen("quiz").ids[self.get_id(instance)].bg_color = (0,1,0,1)
            option_id_list = ["option1", "option2", "option3", "option4"]
            option_id_list.remove(self.get_id(instance))
            for i in range(0,3):
                sm.get_screen("quiz").ids[f"{option_id_list[i]}"].disabled = True
                sm.get_screen("quiz").ids[self.get_id(instance)].bg_color = (0,1,0,1)
        else:
            self.wrong += 1
            for i in range(1,5):
                if sm.get_screen("quiz").ids[f"option{i}"].text == self.answer:
                    sm.get_screen("quiz").ids[f"option{i}"].bg_color = (0,1,0,1)
                else:
                    sm.get_screen("quiz").ids[f"option{i}"].disabled = True
            sm.get_screen("quiz").ids[self.get_id(instance)].bg_color = (1,0,0,1)
            sm.get_screen("quiz").ids[self.get_id(instance)].disabled_color = (1,1,1,1)
    def next_question(self):
        self.select_sign(self.selected_sign)
        for i in range(1, 5):
            sm.get_screen("quiz").ids[f"option{i}"].disabled = False
            sm.get_screen("quiz").ids[f"option{i}"].bg_color = (87/255, 23/255, 216/255, 1)
            sm.get_screen("quiz").ids[f"option{i}"].disabled_color = (1, 1, 1, 0.3)
    def working(self):
        toast("Still in developement")
    def rate_us(self):
        self.r_dialog = MDDialog(type="custom",content_cls=RateUs(),size_hint=(None, None), width=350)
        self.r_dialog.open()
    def feedback(self):
        self.f_dialog = MDDialog(type="custom",content_cls=Feedback(),size_hint=(None, None), width=350)
        self.f_dialog.open()
    def open_url(self, url):
        dialog = MDDialog(
            title="Warning",
            text="You are about to leave the app and visit an external website. Continue?",
            size_hint=(None, None),
            size=(dp(300), dp(200)),
            buttons=[
                MDFlatButton(
                    text="Cancel",
                    on_release=lambda x: dialog.dismiss()
                ),
                MDFlatButton(
                    text="OK",
                    on_release=lambda x: self.open_browser(url)
                )
            ]
        )

        dialog.open()
    def open_browser(self, url):
        global internet_connection
        if internet_connection:
            import webbrowser
            webbrowser.open(url)
        else:
            toast("No internet connection!")
    def ask_chatGPT(self, p):
        for data in chatbot.ask(p):
            self.message = data["message"]
        return self.message
    def send(self):
        global size, halign, value, color, screen15
        screen15 = sm.get_screen("screen15")
        if screen15.text_input != "":
            value = screen15.text_input.text
            if len(value) < 6:
                size = .22
                halign = "center"
            elif len(value) < 11:
                size = .32
                halign = "center"
            elif len(value) < 16:
                size = .45
                halign = "center"
            elif len(value) < 21:
                size = .58
                halign = "center"
            elif len(value) < 26:
                size = .71
                halign = "center"
            else:
                size = .77
                halign = "left"
        if internet_connection:
            if value.upper() == "AI1":
                value = ""
                screen15.chat_list.add_widget(ResponseImage(source="images/AI1.png"))
            elif value.upper() == "AI2":
                value = ""
                screen15.chat_list.add_widget(ResponseImage(source="images/AI2.png"))
            elif value.upper() == "AI3":
                value = ""
                screen15.chat_list.add_widget(ResponseImage(source="images/AI3.png"))
            elif value.upper() == "AI4":
                value = ""
                screen15.chat_list.add_widget(ResponseImage(source="images/AI4.png"))
            else:
                toast("Generating")
                Clock.schedule_once(lambda dt: self.generate_message(value), 1)
        else:
            toast("No internet conection!")
    def generate_message(self, value):
        response = self.ask_chatGPT(value)
        toast("Done")
        screen15.chat_list.add_widget(Command(text=str(value), size_hint_x=size, halign=halign, color=(0, 0, 0, 1)))
        screen15.chat_list.add_widget(Response(text=str(response), size_hint_x=.8, halign="left", color=(0, 0, 0, 1)))
        screen15.text_input.text = ""
    def clear(self):
        screen15.chat_list.clear_widgets()
    def show_task_dialog(self):
        if not self.task_list_dialog:
            self.task_list_dialog = MDDialog(
                title="Create Task",
                type="custom",
                content_cls=DialogContent(),
            )

        self.task_list_dialog.open()
    def on_start(self):
        global internet_connection
        completed_tasks, uncomplete_tasks = db.get_tasks()
        screen5 = self.root.get_screen("screen5")
        if uncomplete_tasks != []:
            for task in uncomplete_tasks:
                add_task = ListItemWithCheckbox(pk=task[0],text=task[1], secondary_text=task[2])
                screen5.ids.container.add_widget(add_task)
        if completed_tasks != []:
            for task in completed_tasks:
                add_task = ListItemWithCheckbox(pk=task[0],text='[s]'+task[1]+'[/s]', secondary_text=task[2])
                add_task.ids.check.active = True
                screen5.ids.container.add_widget(add_task)
        if internet_connection:
            try:
                url = "http://ipinfo.io/json"
                response = urlopen(url)
                data = json.load(response)
                city = (data["city"])
                self.get_weather(city)
            except Exception as e :
                print(e)
                pass
        else:
            toast("No Internet Connection! - M")
    def close_dialog(self, *args):
        self.task_list_dialog.dismiss()
    def add_task(self, task, task_date):
        '''Add task to the list of tasks'''

        created_task = db.create_task(task.text, task_date)
        screen5 = self.root.get_screen("screen5")
        screen5.ids['container'].add_widget(ListItemWithCheckbox(pk=created_task[0], text='[b]'+created_task[1]+'[/b]', secondary_text=created_task[2]))# Here
        task.text = ''
    def get_weather(self, city_name):
        global internet_connection
        if internet_connection:
            try:
                url = f"https://api.openweathermap.org/data/2.5/weather?q={city_name}&appid={self.api_key}"
                x = requests.get(url).json()
                screen6 = self.root.get_screen("screen6")
                if x["cod"] != "404":
                    temperature = str(round(x["main"]["temp"]-273.15))
                    temperature = f"[b]{temperature}[/b]°"
                    humidity = x["main"]["humidity"]
                    weather = x["weather"][0]["main"]
                    id = str(x["weather"][0]["id"])
                    wind_speed = round(x["wind"]["speed"]*18/5)
                    location = x["name"] + ", " + x["sys"]["country"]
                    screen6.ids.temperature.text = str(temperature)
                    #screen4.ids.temperature.text = str(temperature)
                    screen6.ids.weather.text = str(weather)
                    screen6.ids.humidity.text = str(f"{humidity}%")
                    screen6.ids.wind_speed.text = str(f"{wind_speed} km/h")
                    screen6.ids.location.text = str(location)
                    #screen4.ids.location.text = str(location)
                    if id == "800":
                        screen6.ids.weather_image.source = "images/w_sun.png"
                        #screen4.ids.weather_image.source = "images/w_sun.png"
                    elif "200" <= id <= "232":
                        screen6.ids.weather_image.source = "images/w_storm.png"
                        #screen4.ids.weather_image.source = "images/w_storm.png"
                    elif "300" <= id <= "321" and "500"<= id <= "531":
                        screen6.ids.weather_image.source = "images/w_rain.png"
                        #screen4.ids.weather_image.source = "images/w_rain.png"
                    elif "600" <= id <= "622" :
                        screen6.ids.weather_image.source = "images/w_snow.png"
                        #screen4.ids.weather_image.source = "images/w_snow.png"
                    elif "701" <= id <= "781":
                        screen6.ids.weather_image.source = "images/w_haze.png"
                        #screen4.ids.weather_image.source = "images/w_haze.png"
                    elif "801" <= id <=  "804":
                        screen6.ids.weather_image.source = "images/w_clouds.png"
                        #screen4.ids.weather_image.source = "images/w_clouds.png"
                else:
                    toast("City Not Found")
            except Exception as e :
                print(e)
                pass
        else:
            toast("No Internet Connection! - W")
    def search_weather(self):
        screen6 = self.root.get_screen("screen6")
        city_name = screen6.ids.city_name.text
        if city_name != "":
            self.get_weather(city_name)   
    def press(self):
        conn = sqlite3.connect('data.db')
        cursor = conn.cursor()
        cursor.execute('SELECT name FROM users')
        result = cursor.fetchone()
        conn.close()
        if result:
            name = result[0]
        else:
            pass
        screen4 = self.root.get_screen("screen4")
        screen4.ids.home_text.text = str(name)
    def final_score(self):
        if self.correct == 0 and self.wrong == 0:
            sm.current = "screen4"
        else: 
            for i in range(1, 5):
                sm.get_screen("quiz").ids[f"option{i}"].disabled = False
                sm.get_screen("quiz").ids[f"option{i}"].bg_color = (87/255, 23/255, 216/255, 1)
                sm.get_screen("quiz").ids[f"option{i}"].disabled_color = (1, 1, 1, 0.3)
            success_rate = round((self.correct/(self.correct+self.wrong))*100)
            sm.get_screen("screen10").correct.text = f"{self.correct} - Correct!"
            sm.get_screen("screen10").wrong.text = f"{self.wrong} - Wrong!"
            sm.get_screen("screen10").success_rate.text = f"{success_rate}% - Success!"
            sm.current = "screen10"
    def replay(self):
        self.correct = 0
        self.wrong = 0
        sm.current = "screen4"
 

if __name__ == '__main__':
    LabelBase.register(name='Poppins', fn_regular='fonts/r_Poppins.ttf')
    LabelBase.register(name='Poppins-Bold', fn_regular='fonts/r_Poppins-Bold.ttf')
    LabelBase.register(name='Roboto-Bold', fn_regular='fonts/r_Roboto-Bold.ttf')
    LabelBase.register(name='Roboto-Medium', fn_regular='fonts/r_Roboto-Medium.ttf')
    LabelBase.register(name='Roboto', fn_regular='fonts/r_Roboto-Regular.ttf')
    LabelBase.register(name='Graublau', fn_regular='fonts/r_Graublau.ttf')
    LabelBase.register(name='LG', fn_regular='fonts/r_LG.ttf')
    accont = Accont()
    db = Database()
    app = MainApp()
    app.run()
